import os
import zipfile
import matplotlib.pyplot as plt
from transformers import pipeline
import whisper
from pyannote.audio import Pipeline
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# ------------------------
# 1. Transcription
# ------------------------
def transcribe_audio(file_path):
    model = whisper.load_model("base")
    result = model.transcribe(file_path)
    return result["text"]

# ------------------------
# 2. Speaker Diarization
# ------------------------
def diarize_audio(file_path, hf_token):
    diarization_pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization", use_auth_token=hf_token
    )
    diarization = diarization_pipeline(file_path)
    segments = []
    for turn, _, speaker in diarization.itertracks(yield_label=True):
        segments.append((turn.start, turn.end, speaker))
    return segments

# ------------------------
# 3. Sentiment Analysis
# ------------------------
def analyze_sentiment(transcript_by_speaker):
    sentiment_pipeline = pipeline("sentiment-analysis")
    speaker_sentiments = {}
    for speaker, text in transcript_by_speaker.items():
        if text.strip():
            speaker_sentiments[speaker] = sentiment_pipeline(text)
    return speaker_sentiments

def plot_sentiment_timeline(speaker_sentiments, output_file="sentiment_timeline.png"):
    labels, scores = [], []
    for speaker, results in speaker_sentiments.items():
        labels.append(speaker)
        scores.append(results[0]['score'] if results else 0)

    plt.figure(figsize=(6,4))
    plt.bar(labels, scores, color="skyblue")
    plt.title("Meeting Sentiment by Speaker")
    plt.ylabel("Sentiment Score")
    plt.ylim(0, 1)
    plt.savefig(output_file)
    plt.close()
    return output_file

# ------------------------
# 4. Summarization
# ------------------------
def summarize_text(text, length="medium"):
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

    if length == "short":
        max_len, min_len = 50, 15
    elif length == "detailed":
        max_len, min_len = 250, 80
    else:  # medium
        max_len, min_len = 130, 40

    summary = summarizer(text, max_length=max_len, min_length=min_len, do_sample=False)
    return summary[0]['summary_text']

# ------------------------
# 5. Report Generation
# ------------------------
def export_reports(transcript, summary, sentiment,
                   docx_file="Meeting_Report.docx",
                   pdf_file="Meeting_Report.pdf"):

    # DOCX
    doc = Document()
    doc.add_heading("Meeting Report", 0)
    doc.add_heading("Transcript", level=1)
    doc.add_paragraph(transcript)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("Sentiment", level=1)
    for speaker, sent in sentiment.items():
        doc.add_paragraph(f"{speaker}: {sent}")
    doc.save(docx_file)

    # PDF
    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate(pdf_file)
    story = [
        Paragraph("Meeting Report", styles["Heading1"]),
        Paragraph("Transcript", styles["Heading2"]),
        Paragraph(transcript, styles["Normal"]),
        Paragraph("Summary", styles["Heading2"]),
        Paragraph(summary, styles["Normal"]),
        Paragraph("Sentiment", styles["Heading2"]),
        Paragraph(str(sentiment), styles["Normal"]),
    ]
    pdf.build(story)

    return [docx_file, pdf_file]

# ------------------------
# 6. Runner Function
# ------------------------
def run_meeting_pipeline(file_path, hf_token, summary_length="medium"):
    print("🎙️ Step 1: Transcribing...")
    transcript = transcribe_audio(file_path)

    print("🗣️ Step 2: Speaker Diarization...")
    diarization = diarize_audio(file_path, hf_token)

    # Simplified: just assign transcript to one speaker
    transcript_by_speaker = {"SPEAKER_00": transcript}

    print("😊 Step 3: Sentiment Analysis...")
    sentiment = analyze_sentiment(transcript_by_speaker)

    print("📊 Step 4: Sentiment Timeline...")
    timeline_file = plot_sentiment_timeline(sentiment)

    print("📝 Step 5: Summarization...")
    summary = summarize_text(transcript, summary_length)

    print("📄 Step 6: Exporting Reports...")
    reports = export_reports(transcript, summary, sentiment)

    print("📦 Step 7: Bundling...")
    zip_filename = "Meeting_Summary_Package.zip"
    with zipfile.ZipFile(zip_filename, "w") as zipf:
        for r in reports:
            zipf.write(r)
        zipf.write(timeline_file)

    print("✅ Pipeline completed successfully!")
    return {
        "transcript": transcript,
        "summary": summary,
        "sentiment": sentiment,
        "reports": reports,
        "zip": zip_filename
    }

# ------------------------
# Run if executed directly
# ------------------------
if __name__ == "__main__":
    # Load token from environment instead of hardcoding
    hf_token = os.getenv("HF_TOKEN")

    if not hf_token:
        raise ValueError("⚠️ Missing Hugging Face token. Please set HF_TOKEN environment variable.")

    output = run_meeting_pipeline("multi_speaker_audio.wav", hf_token, summary_length="detailed")

    print("\n=== FINAL OUTPUT ===")
    print("Summary:", output["summary"])
    print("Sentiment:", output["sentiment"])
    print("ZIP File:", output["zip"])
